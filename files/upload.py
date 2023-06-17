from docx import Document
from files.workbook import Workbook
import os

def create_document(workbook: Workbook):
  document = Document()
  document.add_heading('Workbook: '+str(workbook.title), 0)

  # loop through passage groups:
  for passage_group in workbook.passage_groups:
    document.add_heading('Passage 1: '+str(passage_group.topic), level=1)
    document.add_paragraph(passage_group.passage)
    document.add_paragraph('')
    document.add_paragraph(passage_group.questions)
    document.add_paragraph('')
    document.add_paragraph('Answers: '+passage_group.answers)

    document.add_page_break()

  script_dir = os.path.dirname(__file__) #<-- absolute dir the script is in
  rel_path_docx = "text/saved-docs/"+workbook.title+".docx"
  # abs_file_path = os.path.join(script_dir, rel_path)
  saved_docx_path = os.path.join(script_dir, rel_path_docx)

  document.save(saved_docx_path)
  print("Workbook Created!")


